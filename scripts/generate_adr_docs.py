from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_run_with_bold(para, text):
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        if i % 2 == 1:
            run.bold = True


def setup_doc_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)


def is_table_line(line):
    return line.strip().startswith('|') and line.strip().endswith('|')


def is_separator_line(line):
    return bool(re.match(r'^\|[\s\-|:]+\|$', line.strip()))


def parse_table_lines(lines, start_idx):
    table_lines = []
    i = start_idx
    while i < len(lines) and (is_table_line(lines[i]) or lines[i].strip() == ''):
        if is_table_line(lines[i]) and not is_separator_line(lines[i]):
            cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
            table_lines.append(cells)
        i += 1
    return table_lines, i


def add_table_to_doc(doc, table_data):
    if not table_data:
        return
    cols = max(len(r) for r in table_data)
    rows = len(table_data)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for ri, row_data in enumerate(table_data):
        row = table.rows[ri]
        for ci in range(cols):
            cell_text = row_data[ci] if ci < len(row_data) else ''
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri == 0:
                set_cell_bg(cell, '2C2C2C')
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cell_text)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            else:
                if ri % 2 == 0:
                    set_cell_bg(cell, 'F5F5F5')
                p = cell.paragraphs[0]
                clean = cell_text.replace('`', '')
                add_run_with_bold(p, clean)
                for run in p.runs:
                    run.font.size = Pt(10)

    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(6.0 / cols)


def parse_and_write(doc, lines):
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blok kodu
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                in_code_block = False
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x2D, 0x6A, 0x4F)
                    pPr = p._p.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'F0F0F0')
                    pPr.append(shd)
                i += 1
                continue

        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue

        # Separator
        if re.match(r'^-{3,}$', stripped):
            add_horizontal_line(doc)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('> '):
            bq_lines = [stripped[2:].strip()]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith('> '):
                bq_lines.append(lines[j].strip()[2:].strip())
                j += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run('\n'.join(bq_lines))
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.size = Pt(10)
            i = j
            continue

        # Tabela
        if is_table_line(stripped):
            table_data, next_i = parse_table_lines(lines, i)
            if table_data:
                add_table_to_doc(doc, table_data)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
            i = next_i
            continue

        # Nagłówki
        m = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            if level == 1:
                p = doc.add_heading('', level=0)
                p.clear()
                run = p.add_run(text)
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(12)
            elif level == 2:
                p = doc.add_heading('', level=1)
                p.clear()
                run = p.add_run(text)
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(6)
            elif level == 3:
                p = doc.add_heading('', level=2)
                p.clear()
                run = p.add_run(text)
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.italic = True
            i += 1
            continue

        # Lista punktowana
        m_bullet = re.match(r'^[-*]\s+(.*)', stripped)
        if m_bullet:
            text = m_bullet.group(1)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(2)
            add_run_with_bold(p, text)
            for run in p.runs:
                run.font.size = Pt(11)
            i += 1
            continue

        # Lista numerowana
        m_num = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if m_num:
            text = m_num.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(2)
            add_run_with_bold(p, text)
            for run in p.runs:
                run.font.size = Pt(11)
            i += 1
            continue

        # Pusty wiersz
        if not stripped:
            i += 1
            continue

        # Zwykly akapit
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_run_with_bold(p, stripped)
        for run in p.runs:
            run.font.size = Pt(11)
        i += 1


def generate_doc(md_path, out_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()
    setup_doc_styles(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    parse_and_write(doc, lines)
    doc.save(out_path)
    print(f"  OK: {os.path.basename(out_path)}")


# MAIN
base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
docs_dir = os.path.join(base_dir, 'docs')
output_dir = os.path.join(docs_dir, 'word')
os.makedirs(output_dir, exist_ok=True)

files = [
    ('ADR-TYDZIEN-1.md',        'ADR-Tydzien-1.docx'),
    ('ADR-TYDZIEN-2.md',        'ADR-Tydzien-2.docx'),
    ('ADR-TYDZIEN-3-NIECALY.md','ADR-Tydzien-3.docx'),
    ('ADR-TYDZIEN-4.md',        'ADR-Tydzien-4.docx'),
]

print("Generuje dokumenty Word...")
for md_name, docx_name in files:
    md_path = os.path.join(docs_dir, md_name)
    out_path = os.path.join(output_dir, docx_name)
    generate_doc(md_path, out_path)

print(f"\nGotowe! Zapisane w: {output_dir}")
