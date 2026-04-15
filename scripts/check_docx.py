from docx import Document
from pathlib import Path

doc = Document(Path(__file__).parent.parent / "ebook-bezp.docx")

print("=== WSZYSTKIE UNIKALNE STYLE ===")
styles = set()
for para in doc.paragraphs:
    if para.text.strip():
        styles.add(para.style.name)
print(styles)

print("\n=== PEŁNA STRUKTURA DOKUMENTU ===")
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        print(f"[{para.style.name}] {text[:80]}")
